"""Text extractor service — pulls raw text out of resume files.

PDFs via pdfplumber (preserves layout, fast for text-based PDFs).
DOCX via python-docx (walks paragraphs + tables in order).

Deterministic, no LLM. The text returned here is the *only* input to the
LLM parser in the next pipeline stage, so this module must:

- Reject unknown extensions early (ValueError).
- Reject empty extractions (RuntimeError — likely scanned/image PDF).
- Reject oversized files before opening them (ValueError).
- Cap output length so a giant resume can't blow the LLM context window.

All file-type discrimination is done by extension (file_type arg). We
trust the API layer to validate MIME at upload time, but re-check here
defensively.
"""
from __future__ import annotations

from pathlib import Path

import docx
import pdfplumber

from app.core.logging import get_logger

log = get_logger(__name__)

# Hard limits. These mirror Settings.max_upload_bytes but are also enforced
# at the extractor level so the service is safe to call from anywhere.
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_LENGTH = 100_000  # chars — well within any modern LLM context window

SUPPORTED_TYPES: tuple[str, ...] = ("pdf", "docx")


def extract_text(file_path: str | Path, file_type: str) -> str:
    """Extract raw text from a resume file.

    Args:
        file_path: Path to the file on disk (must exist).
        file_type: Either ``"pdf"`` or ``"docx"`` (case-insensitive).

    Returns:
        The extracted text, joined with newlines. Whitespace is preserved
        as-is (the LLM will re-tokenize).

    Raises:
        ValueError: Unknown file_type, file missing, or file too large.
        RuntimeError: Extraction succeeded but produced empty text
            (likely a scanned PDF or corrupt DOCX).
    """
    path = Path(file_path)
    file_type = (file_type or "").lower().strip()

    if file_type not in SUPPORTED_TYPES:
        raise ValueError(
            f"unsupported file_type '{file_type}'; expected one of {SUPPORTED_TYPES}"
        )
    if not path.exists():
        raise ValueError(f"file not found: {path}")

    size = path.stat().st_size
    if size > MAX_FILE_SIZE:
        raise ValueError(
            f"file too large: {size} bytes (max {MAX_FILE_SIZE})"
        )
    if size == 0:
        raise ValueError("file is empty (0 bytes)")

    if file_type == "pdf":
        text = _extract_pdf(path)
    else:  # docx
        text = _extract_docx(path)

    text = text.strip()
    if not text:
        raise RuntimeError(
            f"extracted text is empty ({file_type} may be scanned/image-based or corrupt)"
        )
    if len(text) > MAX_TEXT_LENGTH:
        log.warning(
            "text_truncated",
            path=str(path),
            file_type=file_type,
            original_len=len(text),
            cap=MAX_TEXT_LENGTH,
        )
        text = text[:MAX_TEXT_LENGTH]

    return text


# ── Private helpers ────────────────────────────────────────────────


def _extract_pdf(path: Path) -> str:
    """Pull text out of every page of a PDF, joined by blank lines.

    Empty pages (no extractable text) are skipped — that way a single
    blank cover page doesn't poison the whole document.
    """
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                parts.append(page_text)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    """Pull text out of a DOCX, walking paragraphs in document order.

    Tables are included too — many resumes use a 2-column table for
    headers/sidebars that we don't want to miss.
    """
    doc = docx.Document(str(path))
    parts: list[str] = []

    # Walk the body in document order. python-docx exposes paragraphs +
    # tables separately, so we iterate both and pick whichever comes
    # next by document.element order.
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag.endswith("}p") and child in para_map:
            text = para_map[child].text or ""
            if text.strip():
                parts.append(text)
        elif child.tag.endswith("}tbl") and child in table_map:
            for row in table_map[child].rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

    # Fallback: if document order didn't yield anything (e.g., all content
    # is in headers/footers), grab paragraphs anyway.
    if not parts:
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)

    return "\n".join(parts)
