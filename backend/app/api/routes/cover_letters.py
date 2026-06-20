"""Phase 9A — Cover Letter CRUD + generation + export routes.

Endpoints:
- POST   /api/cover-letters/generate   generate a new draft (deterministic + optional LLM)
- GET    /api/cover-letters            list the user's cover letters
- GET    /api/cover-letters/{id}       get one
- PATCH  /api/cover-letters/{id}       edit (content, subject, tone, status)
- DELETE /api/cover-letters/{id}       delete
- POST   /api/cover-letters/{id}/rescore
- POST   /api/cover-letters/{id}/export?format=pdf|docx
- GET    /api/cover-letters/{id}/exports

All routes verify ownership via the profile_id chain (cover_letter
→ profile → user_id == current_user.id). Same pattern as CV routes.
"""
from __future__ import annotations

import re
from datetime import datetime, timezone
from typing import Any

import structlog
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session, joinedload

from app.db.session import get_db
from app.models.models import CoverLetter, Export, Job, Profile, User
from app.schemas.schemas import CoverLetterIn, CoverLetterOut, CoverLetterPatchIn, ExportOut
from app.services.cover_letter_generator import (
    generate_cover_letter_deterministic,
    enhance_cover_letter,
    score_cover_letter,
)

from .jobs import get_current_user

log = structlog.get_logger(__name__)
router = APIRouter(prefix="/cover-letters", tags=["cover-letters"])


# ── Helpers ─────────────────────────────────────────────────────────


# L9 fix (Phase 9 review): hard cap to keep cover-letter content
# rows bounded. 20_000 chars is ~3-4 pages of dense prose — well
# above the realistic upper end of a recruiter-friendly letter.
_MAX_CONTENT_CHARS = 20_000

# H4 fix (Phase 9 review): shared safe-title helper used by both the
# CV and cover-letter exporters. Keeps Content-Disposition filenames
# well-formed even when the source subject contains path separators,
# quotes, or control chars. Same allowlist used by cvs.py:export_cv
# so behavior is identical across both exporters.
_SAFE_TITLE_RE = re.compile(r"[^A-Za-z0-9._\- ]+")
_SAFE_TITLE_COLLAPSE_RE = re.compile(r"\s+")


def _safe_filename_title(raw: str | None, fallback: str, max_len: int = 80) -> str:
    """Sanitize an arbitrary string for use in ``Content-Disposition: filename=``.

    Strips any chars outside ``[A-Za-z0-9._- ]`` and collapses runs of
    whitespace. Falls back to ``fallback`` if the result is empty.
    Truncates to ``max_len`` characters.
    """
    if not raw:
        return fallback[:max_len]
    cleaned = _SAFE_TITLE_RE.sub("", raw).strip()
    cleaned = _SAFE_TITLE_COLLAPSE_RE.sub("-", cleaned).strip("-")
    if not cleaned:
        return fallback[:max_len]
    return cleaned[:max_len]


# L5 fix (Phase 9 review): single-query ownership check via joinedload.
# Avoids the two round-trips (CoverLetter.get → Profile.get) the old
# helper did, where one suffices with a JOIN.
def _get_owned_cover_letter(
    db: Session, cover_letter_id: str, user: User
) -> CoverLetter:
    cl = (
        db.query(CoverLetter)
        .options(joinedload(CoverLetter.profile))
        .filter(CoverLetter.id == cover_letter_id)
        .first()
    )
    if cl is None:
        raise HTTPException(404, f"cover letter {cover_letter_id} not found")
    profile = cl.profile  # eager-loaded, no extra round trip
    if profile is None or profile.user_id != user.id:
        raise HTTPException(403, "not your cover letter")
    return cl


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _serialize_cover_letter(cl: CoverLetter) -> dict[str, Any]:
    """Convert ORM → dict matching CoverLetterOut shape."""
    # L8 fix (Phase 9 review): drop the defensive ``dict(...)`` copy —
    # the ORM already returns a fresh dict per attribute access, and
    # the Pydantic response_model will validate on its own.
    return {
        "id": cl.id,
        "job_id": cl.job_id,
        "profile_id": cl.profile_id,
        "cv_draft_id": cl.cv_draft_id,
        "tone": cl.tone,
        "subject": cl.subject,
        "content": cl.content,
        "personalization_points": list(cl.personalization_points or []),
        "job_keywords_used": list(cl.job_keywords_used or []),
        "score": cl.score or 0.0,
        "score_breakdown_json": cl.score_breakdown_json or {},
        "status": cl.status,
        "created_at": cl.created_at,
        "updated_at": cl.updated_at,
    }


# ── Generate (POST /generate) ────────────────────────────────────────


@router.post("/generate", response_model=CoverLetterOut)
async def generate_cover_letter(
    payload: dict[str, Any],
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Generate a new cover letter for a job.

    Payload:
        job_id: required
        profile_id: required
        cv_draft_id: optional
        tone: one of professional|confident|friendly|concise|formal
        use_llm: bool (default True) — set False for deterministic-only

    Creates a new CoverLetter row with status='draft' and the
    generated content. Re-running for the same job returns the
    latest draft (deterministic by latest tone).
    """
    job_id = payload.get("job_id")
    profile_id = payload.get("profile_id")
    if not job_id or not profile_id:
        raise HTTPException(400, "job_id and profile_id are required")

    tone = payload.get("tone") or "professional"
    if tone not in ("professional", "confident", "friendly", "concise", "formal"):
        raise HTTPException(400, f"invalid tone: {tone}")

    use_llm = bool(payload.get("use_llm", True))

    # ── ownership: verify profile + job belong to user ────────────
    profile = db.get(Profile, profile_id)
    if profile is None:
        raise HTTPException(404, f"profile {profile_id} not found")
    if profile.user_id != user.id:
        raise HTTPException(403, "not your profile")

    job = db.get(Job, job_id)
    if job is None:
        raise HTTPException(404, f"job {job_id} not found")
    if job.user_id != user.id:
        raise HTTPException(403, "not your job")

    # ── deterministic first (always works) ────────────────────────
    profile_dict = dict(profile.base_profile_json or {})
    job_dict = {
        "title": job.title,
        "company": job.company,
        "description": job.raw_description,
        "job_analysis_json": dict(job.job_analysis_json or {}),
        "ats_keywords": list((job.ats_keywords_json or {}).get("keywords") or []),
    }

    draft = generate_cover_letter_deterministic(profile_dict, job_dict, tone=tone)

    # ── optional LLM pass ─────────────────────────────────────────
    if use_llm:
        try:
            draft = await enhance_cover_letter(draft, profile_dict, job_dict)
        except (RuntimeError, TimeoutError, ConnectionError, OSError) as exc:
            # M2 fix (Phase 9 review): narrow from ``Exception`` so
            # genuine bugs surface instead of silently degrading to
            # deterministic. Mirrors the narrow in
            # ``cover_letter_generator.enhance_cover_letter``.
            log.warning(
                "cover_letter_generate_llm_failed",
                error=str(exc)[:200],
                job_id=job_id,
            )
            # Fall back to deterministic — still works.

    # ── score ─────────────────────────────────────────────────────
    required_keywords: list[str] = list(
        (job.job_analysis_json or {}).get("required_skills") or []
    ) or list((job.ats_keywords_json or {}).get("keywords") or [])
    score_result = score_cover_letter(draft.body, required_keywords)

    # ── persist ───────────────────────────────────────────────────
    cv_draft_id = payload.get("cv_draft_id") or None
    cl = CoverLetter(
        job_id=job_id,
        profile_id=profile_id,
        cv_draft_id=cv_draft_id,
        tone=tone,
        subject=draft.subject,
        content=draft.body,
        personalization_points=draft.personalization_points,
        job_keywords_used=draft.job_keywords_used,
        score=score_result.overall,
        score_breakdown_json=score_result.to_breakdown(),
        status="draft",
    )
    db.add(cl)
    db.commit()
    db.refresh(cl)

    log.info(
        "cover_letter_generated",
        cover_letter_id=cl.id,
        job_id=job_id,
        tone=tone,
        source=draft.source,
        score=score_result.overall,
    )

    return _serialize_cover_letter(cl)


# ── List (GET /) ────────────────────────────────────────────────────


@router.get("", response_model=list[CoverLetterOut])
def list_cover_letters(
    job_id: str | None = Query(None, description="Filter by job"),
    status: str | None = Query(None, description="Filter by status"),
    limit: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    """List the user's cover letters, newest first."""
    q = (
        db.query(CoverLetter)
        .join(Profile, CoverLetter.profile_id == Profile.id)
        .filter(Profile.user_id == user.id)
    )
    if job_id:
        q = q.filter(CoverLetter.job_id == job_id)
    if status:
        q = q.filter(CoverLetter.status == status)
    rows = q.order_by(CoverLetter.updated_at.desc()).limit(limit).all()
    return [_serialize_cover_letter(r) for r in rows]


# ── Get one (GET /{id}) ─────────────────────────────────────────────


@router.get("/{cover_letter_id}", response_model=CoverLetterOut)
def get_cover_letter(
    cover_letter_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    cl = _get_owned_cover_letter(db, cover_letter_id, user)
    return _serialize_cover_letter(cl)


# ── Patch (PATCH /{id}) ─────────────────────────────────────────────


@router.patch("/{cover_letter_id}", response_model=CoverLetterOut)
def patch_cover_letter(
    cover_letter_id: str,
    # H5 fix (Phase 9 review): switch from ``dict[str, Any]`` to the
    # existing ``CoverLetterIn`` Pydantic model so ``tone`` and
    # ``status`` enum validation happen at the framework boundary
    # instead of via hand-rolled `if value not in (...)` checks that
    # drift if the enum grows.
    payload: CoverLetterPatchIn,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Edit a cover letter. Allowed: content, subject, tone, status,
    personalization_points, job_keywords_used.

    Editing `content` triggers a re-score automatically (same
    pattern as CV's _score_and_persist).
    """
    cl = _get_owned_cover_letter(db, cover_letter_id, user)

    if payload.tone is not None:
        cl.tone = payload.tone
    if payload.status is not None:
        cl.status = payload.status
    if payload.subject is not None:
        cl.subject = str(payload.subject)[:500]
    if payload.personalization_points is not None:
        cl.personalization_points = list(payload.personalization_points)
    if payload.job_keywords_used is not None:
        cl.job_keywords_used = list(payload.job_keywords_used)
    if payload.content is not None:
        # L9 fix: enforce a hard cap on content length to keep DB
        # rows bounded. A malicious or accidental huge PATCH would
        # otherwise balloon the row (TEXT but still costly).
        if len(payload.content) > _MAX_CONTENT_CHARS:
            raise HTTPException(
                400,
                f"content too long ({len(payload.content)} chars > "
                f"{_MAX_CONTENT_CHARS} max)",
            )
        cl.content = payload.content

    # Re-score if content or keywords changed
    if payload.content is not None or payload.job_keywords_used is not None:
        job = db.get(Job, cl.job_id)
        required_keywords: list[str] = list(
            (job.job_analysis_json or {}).get("required_skills") or []
        ) if job else []
        score_result = score_cover_letter(cl.content, required_keywords)
        cl.score = score_result.overall
        cl.score_breakdown_json = score_result.to_breakdown()

    cl.updated_at = _utcnow()
    db.commit()
    db.refresh(cl)

    log.info("cover_letter_updated", cover_letter_id=cl.id, fields=list(payload.model_dump(exclude_unset=True).keys()))
    return _serialize_cover_letter(cl)


# ── Rescore (POST /{id}/rescore) ───────────────────────────────────


@router.post("/{cover_letter_id}/rescore", response_model=CoverLetterOut)
def rescore_cover_letter(
    cover_letter_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Recompute the score (e.g. after a job is re-analyzed)."""
    cl = _get_owned_cover_letter(db, cover_letter_id, user)
    job = db.get(Job, cl.job_id)
    required_keywords: list[str] = list(
        (job.job_analysis_json or {}).get("required_skills") or []
    ) if job else []
    score_result = score_cover_letter(cl.content, required_keywords)
    cl.score = score_result.overall
    cl.score_breakdown_json = score_result.to_breakdown()
    cl.updated_at = _utcnow()
    db.commit()
    db.refresh(cl)
    # M4 fix (Phase 9 review): audit log for completeness — every
    # mutating endpoint should leave a structured event trail.
    log.info(
        "cover_letter_rescored",
        cover_letter_id=cl.id,
        new_score=cl.score,
    )
    return _serialize_cover_letter(cl)


# ── Delete (DELETE /{id}) ───────────────────────────────────────────


@router.delete("/{cover_letter_id}", status_code=204)
def delete_cover_letter(
    cover_letter_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> Response:
    cl = _get_owned_cover_letter(db, cover_letter_id, user)
    db.delete(cl)
    db.commit()
    # M4 fix (Phase 9 review): audit log for completeness — every
    # mutating endpoint should leave a structured event trail.
    log.info("cover_letter_deleted", cover_letter_id=cover_letter_id)
    return Response(status_code=204)


# ── Export (POST /{id}/export) ─────────────────────────────────────


def _render_cover_letter_html(cl: CoverLetter, profile: Profile) -> str:
    """Render a cover letter as ATS-safe HTML for WeasyPrint."""
    basics = (profile.base_profile_json or {}).get("basics") or {}
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{cl.subject or "Cover Letter"}</title>
<style>
  body {{
    font-family: Arial, Helvetica, sans-serif;
    color: #111;
    max-width: 720px;
    margin: 24px auto;
    padding: 0 16px;
    line-height: 1.55;
    font-size: 14px;
  }}
  .cl-header {{ margin-bottom: 24px; }}
  .cl-from {{ color: #444; font-size: 13px; }}
  .cl-from strong {{ color: #111; }}
  .cl-subject {{ margin: 18px 0; font-weight: 600; }}
  .cl-body {{ white-space: pre-wrap; }}
  .cl-body p {{ margin: 0 0 12px; }}
</style>
</head>
<body>
  <div class="cl-header">
    <div class="cl-from"><strong>{basics.get("name", "")}</strong></div>
    <div class="cl-from">{basics.get("email", "")} &middot; {basics.get("phone", "")} &middot; {basics.get("location", "")}</div>
  </div>
  <div class="cl-subject">{cl.subject or ""}</div>
  <div class="cl-body">{cl.content}</div>
</body>
</html>
"""


@router.post("/{cover_letter_id}/export", response_class=Response)
def export_cover_letter(
    cover_letter_id: str,
    fmt: str = Query(
        "pdf",
        alias="format",
        pattern="^(pdf|docx)$",
        description="Output format. pdf is wired up; docx is wired up too.",
    ),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Export a cover letter as PDF (WeasyPrint) or DOCX (python-docx).

    Reuses the ATS-safe CSS from cv_exporter for PDF; for DOCX uses
    python-docx with a minimal styled document.
    """
    import hashlib

    cl = _get_owned_cover_letter(db, cover_letter_id, user)
    profile = db.get(Profile, cl.profile_id)

    if fmt == "pdf":
        from app.services.cv_exporter import export_cv_to_pdf, pdf_metadata

        html = _render_cover_letter_html(cl, profile)
        try:
            pdf_bytes = export_cv_to_pdf(html)
        except (RuntimeError, OSError, ValueError) as exc:
            # M2 fix (Phase 9 review): narrow from ``Exception`` so a
            # genuine bug (e.g. AttributeError on a refactor) surfaces
            # instead of being silently swallowed as "PDF failed".
            log.error(
                "cover_letter_pdf_export_failed",
                cover_letter_id=cover_letter_id,
                error=str(exc)[:200],
            )
            try:
                db.add(
                    Export(
                        user_id=user.id,
                        entity_type="cover_letter",
                        entity_id=cover_letter_id,
                        cover_letter_id=cover_letter_id,
                        file_type="failed",
                        file_path=f"on-demand://failed/{cover_letter_id}",
                        file_size=0,
                    )
                )
                db.commit()
            except (RuntimeError, OSError, ValueError):
                db.rollback()
            raise HTTPException(
                500,
                "Cover letter PDF generation failed — please retry",
            )

        meta = pdf_metadata(pdf_bytes)
        if not meta["is_valid"]:
            raise HTTPException(500, "Generated PDF failed validation")

        content_hash = hashlib.sha256(pdf_bytes).hexdigest()
        # H4 fix (Phase 9 review): use the shared safe-title helper
        # so subjects containing quotes, backslashes, or control
        # chars can't break the Content-Disposition header.
        safe_subject = _safe_filename_title(cl.subject, fallback="cover_letter")
        file_name = f"{safe_subject}_v{cl.id[:8]}.pdf"

        exp = Export(
            user_id=user.id,
            entity_type="cover_letter",
            entity_id=cover_letter_id,
            cover_letter_id=cover_letter_id,
            file_type="pdf",
            file_path=f"on-demand://{cover_letter_id}/{file_name}",
            file_size=meta["size"],
            sha256=content_hash,
        )
        db.add(exp)
        db.flush()
        exp.file_path = f"on-demand://{exp.id}/{file_name}"
        db.commit()

        log.info(
            "cover_letter_exported",
            cover_letter_id=cover_letter_id,
            export_id=exp.id,
            pdf_size=meta["size"],
        )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{file_name}"',
                "X-Cover-Letter-Export-Id": exp.id,
                "X-Cover-Letter-Export-Size": str(meta["size"]),
            },
        )

    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO

        doc = Document()
        # Apply a clean default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        basics = (profile.base_profile_json or {}).get("basics") or {}
        # Header
        header = doc.add_paragraph()
        run = header.add_run(basics.get("name", ""))
        run.bold = True
        contact = doc.add_paragraph()
        contact.add_run(
            f"{basics.get('email', '')} \u00b7 {basics.get('phone', '')} \u00b7 {basics.get('location', '')}"
        )
        doc.add_paragraph()  # spacer
        # Subject
        subj_p = doc.add_paragraph()
        subj_run = subj_p.add_run(cl.subject or "")
        subj_run.bold = True
        doc.add_paragraph()  # spacer
        # Body — one paragraph per \n\n
        for para in (cl.content or "").split("\n\n"):
            p = doc.add_paragraph()
            for i, line in enumerate(para.split("\n")):
                if i > 0:
                    p.add_run().add_break()
                p.add_run(line)

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        content_hash = hashlib.sha256(docx_bytes).hexdigest()
        # H4 fix: same safe-title path as the PDF branch above.
        safe_subject = _safe_filename_title(cl.subject, fallback="cover_letter")
        file_name = f"{safe_subject}_v{cl.id[:8]}.docx"

        exp = Export(
            user_id=user.id,
            entity_type="cover_letter",
            entity_id=cover_letter_id,
            cover_letter_id=cover_letter_id,
            file_type="docx",
            file_path=f"on-demand://{cover_letter_id}/{file_name}",
            file_size=len(docx_bytes),
            sha256=content_hash,
        )
        db.add(exp)
        db.flush()
        exp.file_path = f"on-demand://{exp.id}/{file_name}"
        db.commit()

        log.info(
            "cover_letter_exported_docx",
            cover_letter_id=cover_letter_id,
            export_id=exp.id,
            docx_size=len(docx_bytes),
        )

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{file_name}"',
                "X-Cover-Letter-Export-Id": exp.id,
                "X-Cover-Letter-Export-Size": str(len(docx_bytes)),
            },
        )

    # fmt already validated by Query pattern, but defensive:
    raise HTTPException(400, f"unsupported format: {fmt}")


@router.get("/{cover_letter_id}/exports", response_model=list[ExportOut])
def list_cover_letter_exports(
    cover_letter_id: str,
    limit: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> list[ExportOut]:
    cl = _get_owned_cover_letter(db, cover_letter_id, user)
    rows = (
        db.query(Export)
        .filter(
            Export.cover_letter_id == cover_letter_id,
            Export.user_id == user.id,
        )
        .order_by(Export.created_at.desc())
        .limit(limit)
        .all()
    )
    # M1 fix (Phase 9 review): use ExportOut for typed validation +
    # OpenAPI schema. Previously returned ``list[dict]`` which bypassed
    # Pydantic and hid the schema from the FE.
    return [ExportOut.model_validate(r) for r in rows]